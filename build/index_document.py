"""
Оркестратор индексации PDF и DOCX.

КРИТИЧНО ИСПРАВЛЕНО (P0): Имя файла FAISS теперь faiss_index.bin (совпадает с index_writer).
                          Раньше из-за несовпадения имён инкрементальный апдейт молча
                          переписывал индекс коллекции только чанками последнего файла.
КРИТИЧНО ИСПРАВЛЕНО (P1): extract_text_from_docx теперь обходит body по порядку и
                          извлекает ТАБЛИЦЫ (раньше doc.paragraphs пропускал их полностью).
ОПТИМИЗИРОВАНО       (P2): reconstruct_n вместо цикла reconstruct — в 10-50x быстрее
                          при больших индексах.
ОБЪЕДИНЕНО           (P3): Коллекция orders маппится в индекс gost с doc_type="order".
                          Нормативная база (ГОСТы + приказы) живёт в одном индексе,
                          но фильтруется по metadata.doc_type при ретриве.
"""

import os
import sys
import json
import time
import psutil
import faiss
import numpy as np
from pathlib import Path

# =====================================================================
# ЖЕСТКОЕ ОГРАНИЧЕНИЕ НА ВЫЧИСЛИТЕЛЬНЫЕ БИБЛИОТЕКИ (однопоточный режим)
# =====================================================================
os.environ["OMP_NUM_THREADS"] = "1"
os.environ["OPENBLAS_NUM_THREADS"] = "1"
os.environ["MKL_NUM_THREADS"] = "1"
os.environ["NUMEXPR_NUM_THREADS"] = "1"
os.environ["VECLIB_MAXIMUM_THREADS"] = "1"

# =====================================================================
# ОГРАНИЧЕНИЕ ОС: Выделяем процессу строго 2 ядра (0 и 1)
# =====================================================================
try:
    process = psutil.Process(os.getpid())
    process.cpu_affinity([0, 1])
    print("🎯 Системное ограничение: Скрипт изолирован на Ядрах 0 и 1.", flush=True)
except Exception as e:
    print(f"⚠️ Не удалось выставить cpu_affinity: {e}", flush=True)

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from config import PROCESSED_DIR, MIN_CHUNK_SIZE
from .chunking import chunk_text, chunk_tables
from .embedding import encode_chunks
from .index_writer import save_all
from .utils import (
    detect_db, extract_text_from_pdf, extract_tables_from_pdf, file_content_hash
)

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("❌ Библиотека python-docx не найдена. Установи её: pip install python-docx")
    sys.exit(1)

# Пауза в секундах между обработкой новых файлов
PAUSE_BETWEEN_FILES = 5.0

# Карта коллекций: Имя папки в data/raw -> Имя результирующей папки индекса (СТРОГО ЛАТИНИЦА)
# ВАЖНО: "orders" маппится в "gost" — нормативные приказы живут в одном индексе с ГОСТами,
# но различаются по metadata.doc_type ("order" vs "gost").
COLLECTION_MAP = {
    "gost": "gost",
    "orders": "gost",                # ← объединяем приказы с ГОСТами
    "manuals": "manuals",
    "projects_docx_clean": "projects_docx_clean",
    "general": "general",
}


def extract_text_from_docx(docx_path: str) -> dict:
    """
    Умное извлечение текста из DOCX с разбиением на условные страницы по абзацам.

    ИСПРАВЛЕНО: теперь обходит body в порядке появления элементов (w:p и w:tbl вперемешку),
    извлекая и параграфы, и таблицы. Раньше doc.paragraphs пропускал таблицы целиком,
    что критично для projects_docx_clean и manuals.
    """
    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"   ❌ Ошибка открытия DOCX файла {docx_path}: {e}")
        return {}

    pages_data = {}
    current_page = 1
    buffer = []

    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn('w:p'):
            # Параграф — собираем текст из всех w:t внутри
            text = "".join(t.text or "" for t in child.iter(qn('w:t'))).strip()
            if text:
                buffer.append(text)
        elif child.tag == qn('w:tbl'):
            # Таблица — пересобираем в markdown-подобное представление
            rows = []
            for row in child.iter(qn('w:tr')):
                cells = []
                for cell in row.iter(qn('w:tc')):
                    cell_text = " ".join(
                        (t.text or "") for t in cell.iter(qn('w:t'))
                    ).strip()
                    cells.append(cell_text)
                rows.append(" | ".join(cells))
            if rows:
                buffer.append("\n".join(rows))

        # Сбрасываем буфер в "страницу" каждые 15 элементов
        if len(buffer) >= 15:
            pages_data[str(current_page)] = "\n\n".join(buffer)
            buffer = []
            current_page += 1

    if buffer:
        pages_data[str(current_page)] = "\n\n".join(buffer)

    return pages_data


def extract_tables_from_docx(docx_path: str) -> list:
    """
    Отдельное извлечение таблиц из DOCX (для метаданных rows/cols).
    Возвращает список в том же формате, что и extract_tables_from_pdf.
    NOTE: текст таблиц уже включён в extract_text_from_docx, эта функция нужна
    только если хочешь дублировать таблицы как отдельные чанки типа "table".
    По умолчанию не вызывается — см. комментарий в index_file.
    """
    tables = []
    try:
        doc = Document(docx_path)
        for t_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if not rows:
                continue
            tables.append({
                "page": 0,  # Условная страница — для DOCX не используется
                "name": f"Таблица {t_idx + 1}",
                "text": "\n".join(rows),
                "rows": len(table.rows),
                "cols": len(table.columns) if table.rows else 0
            })
    except Exception as e:
        print(f"⚠️ Предупреждение при парсинге таблиц из {docx_path}: {e}")
    return tables


def _detect_doc_type(source_name: str, file_path: str, ext: str) -> tuple:
    """Возвращает (collection, doc_type) на основе имени файла и содержимого."""
    if ext == ".pdf":
        pages_data_preview = extract_text_from_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        pages_data_preview = extract_text_from_docx(file_path)
    else:
        return "general", "general"
    first_text = ""
    for p_key in sorted(pages_data_preview.keys(), key=lambda k: int(k) if k.isdigit() else 0):
        if pages_data_preview[p_key]:
            first_text = pages_data_preview[p_key]
            break
    return detect_db(source_name, first_text[:2000])


def index_file(file_path: str, collection: str = None) -> int:
    """Индексирует файл (PDF или DOCX) с умным инкрементальным обновлением векторов."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Файл не найден: {file_path}")

    ext = Path(file_path).suffix.lower()
    source_name = os.path.basename(file_path)
    fhash = file_content_hash(file_path)

    # 1. Детекция коллекции и типа документа
    if collection is None:
        collection, doc_type = _detect_doc_type(source_name, file_path, ext)
    else:
        # Если коллекция передана явно (через --collection или index_all),
        # всё равно вычисляем doc_type для метаданных
        _, doc_type = _detect_doc_type(source_name, file_path, ext)

    collection_dir = PROCESSED_DIR / "indexes" / collection
    collection_dir.mkdir(parents=True, exist_ok=True)

    chunks_path = collection_dir / "chunks.json"
    # ИСПРАВЛЕНО (P0): имя файла FAISS теперь совпадает с тем, что пишет index_writer.py
    faiss_path = collection_dir / "faiss_index.bin"

    # 2. ПРОВЕРКА ХЭША: Если файл уже есть в базе в неизменном виде — полный скип
    existing_chunks = []
    file_already_indexed = False

    if chunks_path.exists():
        with open(chunks_path, 'r', encoding='utf-8') as f:
            existing_chunks = json.load(f)
        file_already_indexed = any(c["metadata"].get("file_hash") == fhash for c in existing_chunks)

    if file_already_indexed:
        print(f"⚡ [ПРОПУСК] Файл '{source_name}' уже проиндексирован и не изменялся (хэш совпадает).")
        return 0

    # 3. ПАРСИНГ
    print(f"📖 Чтение и парсинг контента: '{source_name}' [doc_type={doc_type}]...")
    if ext == ".pdf":
        pages_data = extract_text_from_pdf(file_path)
        all_tables = extract_tables_from_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        pages_data = extract_text_from_docx(file_path)
        # NOTE: таблицы DOCX уже включены в pages_data через extract_text_from_docx.
        # Если хочешь дублировать их как отдельные чанки типа "table" — раскомментируй:
        # all_tables = extract_tables_from_docx(file_path)
        all_tables = []
    else:
        print(f"⚠️ Неподдерживаемый формат файла: {source_name}")
        return 0

    if not pages_data:
        print(f"⚠️ Файл {source_name} пуст или не распознан.")
        return 0

    tables_by_page = {}
    for tbl in all_tables:
        page = tbl["page"]
        tables_by_page.setdefault(page, []).append(tbl)

    new_file_chunks = []
    current_section = ""
    current_title = ""

    for key in sorted(pages_data.keys(), key=lambda k: int(k) if k.isdigit() else 0):
        text = pages_data[key]
        if not text:
            continue
        page_num = int(key)

        page_chunks, current_section, current_title = chunk_text(
            text, source_name, page_num, fhash,
            current_section, current_title, doc_type=doc_type
        )
        new_file_chunks.extend(page_chunks)

        if page_num in tables_by_page:
            table_chunks = chunk_tables(tables_by_page[page_num], source_name,
                                        current_section, fhash, doc_type=doc_type)
            new_file_chunks.extend(table_chunks)

    # Для DOCX: если extract_tables_from_docx был включён, таблицы привязаны к page=0
    if ext in [".docx", ".doc"] and 0 in tables_by_page:
        table_chunks = chunk_tables(tables_by_page[0], source_name,
                                    current_section, fhash, doc_type=doc_type)
        new_file_chunks.extend(table_chunks)

    new_file_chunks = [c for c in new_file_chunks if len(c["text"]) >= MIN_CHUNK_SIZE]
    if not new_file_chunks:
        print(f"⚠️ Из файла {source_name} не выделено пригодных чанков (минимум {MIN_CHUNK_SIZE} символов).")
        return 0

    # 4. ИНКРЕМЕНТАЛЬНАЯ СБОРКА ИНДЕКСА
    print(f"🧠 Генерируем векторы локально для {len(new_file_chunks)} новых чанков...")
    new_embeddings = encode_chunks(new_file_chunks)

    if chunks_path.exists() and faiss_path.exists():
        # Загружаем старую матрицу FAISS
        old_index = faiss.read_index(str(faiss_path))

        n_total = old_index.ntotal
        dim = old_index.d

        # ИСПРАВЛЕНО (P2): reconstruct_n вместо цикла reconstruct — в 10-50x быстрее
        if n_total > 0:
            old_embeddings = old_index.reconstruct_n(0, n_total).astype(np.float32)
        else:
            old_embeddings = np.zeros((0, dim), dtype=np.float32)

        # Находим индексы чанков, которые принадлежали СТАРОЙ версии этого же файла
        keep_indices = []
        updated_chunks = []

        for idx, c in enumerate(existing_chunks):
            # Сохраняем чанки, не принадлежащие текущему файлу
            if c["metadata"].get("source") != source_name and c["metadata"].get("file_hash") != fhash:
                keep_indices.append(idx)
                updated_chunks.append(c)

        # Формируем очищенную матрицу старых векторов
        if keep_indices and n_total > 0:
            filtered_old_embeddings = old_embeddings[keep_indices]
            final_embeddings = np.vstack([filtered_old_embeddings, new_embeddings])
        else:
            final_embeddings = new_embeddings

        # Объединяем метаданные чанков
        updated_chunks.extend(new_file_chunks)
        final_chunks = updated_chunks
    else:
        # Первая сборка
        final_chunks = new_file_chunks
        final_embeddings = new_embeddings

    # Сохраняем обновленный агрегированный индекс (JSON + FAISS + BM25)
    save_all(final_chunks, final_embeddings, collection_dir)

    print(f"✅ Коллекция '{collection}': успешно добавлено/обновлено {len(new_file_chunks)} чанков из '{source_name}'")
    return len(new_file_chunks)


def index_all():
    raw_root = Path("data/raw")

    indexes_root = PROCESSED_DIR / "indexes"
    indexes_root.mkdir(parents=True, exist_ok=True)

    for coll in COLLECTION_MAP:
        coll_dir = raw_root / coll
        if not coll_dir.exists():
            print(f"⚠️ Папка коллекции raw не найдена и пропущена: {coll_dir}")
            continue

        # Жёсткая фильтрация дубликатов масок через set()
        unique_files = set()
        for ext in ["*.pdf", "*.docx", "*.doc", "*.PDF", "*.DOCX", "*.DOC"]:
            unique_files.update(coll_dir.glob(ext))

        files = sorted(list(unique_files))

        if not files:
            print(f"📁 Папка '{coll}' пуста.")
            continue

        target_collection = COLLECTION_MAP[coll]
        if coll != target_collection:
            print(f"\n📂 Коллекция '{coll}' → индекс '{target_collection}': {len(files)} файлов (ОБЪЕДИНЕНО)")
        else:
            print(f"\n📂 Коллекция '{coll}' → индекс '{target_collection}': {len(files)} файлов")
        for file in files:
            print(f"\nПроверка файла: {file}")
            processed = index_file(str(file), collection=target_collection)

            if processed > 0:
                print(f"💤 Пауза {PAUSE_BETWEEN_FILES} сек. Охлаждаем CPU...")
                time.sleep(PAUSE_BETWEEN_FILES)


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("file", nargs="?")
    parser.add_argument("--collection")
    parser.add_argument("--all", action="store_true")
    args = parser.parse_args()

    if args.all:
        index_all()
    elif args.file:
        index_file(args.file, args.collection)
    else:
        print("Укажите --all или путь к файлу")