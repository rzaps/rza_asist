"""
RZA_Docx_Assembler — Компонент №4 пайплайна.
Собирает сгенерированные разделы в документ .docx
с автонумерацией, оглавлением и базовым форматированием по ГОСТ.
"""

import sys
from pathlib import Path
from typing import Optional

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
except ImportError:
    print("❌ Установи python-docx: pip install python-docx")
    raise

from pipeline._04_orchestrator import PipelineState

# ============================================================
# КОНСТАНТЫ ФОРМАТИРОВАНИЯ
# ============================================================

FONT_NAME = "Times New Roman"
FONT_SIZE_TEXT = Pt(14)     # Основной текст
FONT_SIZE_HEADING1 = Pt(16) # Заголовок раздела (жирный)
FONT_SIZE_HEADING2 = Pt(14) # Подзаголовок
FONT_SIZE_TITLE = Pt(18)    # Титул
FONT_SIZE_SMALL = Pt(10)    # Колонтитулы, примечания

LINE_SPACING = 1.5  # Полуторный межстрочный интервал
MARGIN_LEFT = Cm(3)     # Левое поле (под подшивку)
MARGIN_RIGHT = Cm(1)
MARGIN_TOP = Cm(2)
MARGIN_BOTTOM = Cm(2)

PARAGRAPH_SPACING = Pt(6)  # Отбивка между абзацами
HEADING_SPACING_BEFORE = Pt(18)  # Отбивка перед заголовком


# ============================================================
# НАСТРОЙКА СТИЛЕЙ ДОКУМЕНТА
# ============================================================

def _setup_styles(doc: Document):
    """Настраивает базовые стили документа."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE_TEXT
    style.paragraph_format.line_spacing = LINE_SPACING
    style.paragraph_format.space_after = PARAGRAPH_SPACING

    # Стиль заголовка 1
    h1 = doc.styles["Heading 1"]
    h1.font.name = FONT_NAME
    h1.font.size = FONT_SIZE_HEADING1
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.space_before = HEADING_SPACING_BEFORE
    h1.paragraph_format.space_after = Pt(6)

    # Стиль заголовка 2
    h2 = doc.styles["Heading 2"]
    h2.font.name = FONT_NAME
    h2.font.size = FONT_SIZE_HEADING2
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.space_before = Pt(12)


def _setup_page(doc: Document):
    """Настраивает поля страницы."""
    for section in doc.sections:
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin = MARGIN_LEFT
        section.right_margin = MARGIN_RIGHT


# ============================================================
# ТИТУЛЬНАЯ ИНФОРМАЦИЯ
# ============================================================

def _add_title_page(doc: Document, state: PipelineState):
    """Добавляет титульную информацию в начало документа."""
    card = state.card or {}

    # Заголовок документа
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ПОЯСНИТЕЛЬНАЯ ЗАПИСКА")
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_TITLE
    run.font.bold = True

    # Стадия
    stage = doc.add_paragraph()
    stage.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = stage.add_run("Стадия: ОТР")
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_HEADING1

    doc.add_paragraph()  # Отступ

    # Объект
    object_name = card.get("object_name") or "________________________"
    obj = doc.add_paragraph()
    obj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = obj.add_run(f"Объект: {object_name}")
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_HEADING1

    # Класс напряжения
    v_hv = card.get("voltage_hv") or "___ кВ"
    v_lv = card.get("voltage_lv")
    voltage_str = f"{v_hv}{' / ' + v_lv if v_lv else ''}"
    voltage = doc.add_paragraph()
    voltage.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = voltage.add_run(f"Класс напряжения: {voltage_str}")
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_HEADING1

    doc.add_paragraph()
    doc.add_paragraph()

    # Место и год (пустые строки для заполнения)
    place = doc.add_paragraph()
    place.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = place.add_run("________________________")
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_TEXT

    year = doc.add_paragraph()
    year.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = year.add_run("2026 г.")
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_TEXT

    # Разрыв страницы после титула
    doc.add_page_break()


# ============================================================
# СОДЕРЖАНИЕ
# ============================================================

def _add_table_of_contents(doc: Document, state: PipelineState):
    """Добавляет ручное оглавление."""
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_title.add_run("СОДЕРЖАНИЕ")
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_HEADING1
    run.font.bold = True

    doc.add_paragraph()

    for section in state.sections:
        num = section.get("section_num", "?")
        title = section.get("section_title", "")
        is_fallback = section.get("is_fallback", False)

        p = doc.add_paragraph()
        fallback_mark = " (заглушка)" if is_fallback else ""
        run = p.add_run(f"{num}. {title}{fallback_mark}")
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE_TEXT

    doc.add_page_break()


# ============================================================
# ТЕКСТЫ РАЗДЕЛОВ
# ============================================================

def _add_sections(doc: Document, state: PipelineState):
    """Добавляет все разделы в документ."""
    for section in state.sections:
        num = section.get("section_num", "?")
        title = section.get("section_title", "")
        text = section.get("text", "")
        is_fallback = section.get("is_fallback", False)

        # Заголовок раздела
        heading = doc.add_heading(f"{num}. {title}", level=1)

        # Метка заглушки
        if is_fallback:
            warning = doc.add_paragraph()
            run = warning.add_run("⚠️ Раздел-заглушка: данные не найдены в базе знаний.")
            run.font.name = FONT_NAME
            run.font.size = FONT_SIZE_SMALL
            run.font.italic = True
            run.font.color.rgb = RGBColor(180, 0, 0)

        # Тело раздела
        paragraphs = text.split("\n")
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue

            # Подзаголовки (эвристика: короткая строка БЕЗ точки в конце)
            if len(para_text) < 80 and not para_text.endswith(".") and not para_text.endswith(":"):
                # Проверяем, похоже ли на подзаголовок
                is_subheading = (
                    para_text.isupper() or
                    para_text[0].isupper() and "требования" in para_text.lower()
                )
                if is_subheading:
                    h = doc.add_heading(para_text, level=2)
                    continue

            p = doc.add_paragraph()
            run = p.add_run(para_text)
            run.font.name = FONT_NAME
            run.font.size = FONT_SIZE_TEXT

        # Источники (мелким шрифтом)
        sources = section.get("sources", [])
        if sources:
            doc.add_paragraph()
            src_title = doc.add_paragraph()
            run = src_title.add_run("Источники:")
            run.font.name = FONT_NAME
            run.font.size = FONT_SIZE_SMALL
            run.font.italic = True

            for src in sources:
                src_text = (
                    f"[{src.get('num', '?')}] "
                    f"{src.get('source', '?')} | "
                    f"Раздел {src.get('section', '?')} | "
                    f"{src.get('collection', '?')}"
                )
                sp = doc.add_paragraph()
                run = sp.add_run(src_text)
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE_SMALL
                run.font.italic = True

        # Отбивка между разделами
        doc.add_paragraph()


# ============================================================
# ОСНОВНАЯ ФУНКЦИЯ
# ============================================================

def assemble_docx(
    state: PipelineState,
    output_path: Optional[str] = None,
    include_title_page: bool = True,
    include_toc: bool = True,
) -> str:
    """
    Собирает все разделы в документ .docx.

    Аргументы:
        state: PipelineState после завершения пайплайна.
        output_path: путь для сохранения (None = авто).
        include_title_page: добавлять титульный лист.
        include_toc: добавлять содержание.

    Возвращает:
        Путь к сохранённому файлу .docx.
    """
    if not state.sections:
        raise ValueError("PipelineState не содержит разделов. Сначала запустите пайплайн.")

    doc = Document()

    # Настройка стилей и страницы
    _setup_styles(doc)
    _setup_page(doc)

    # Титульный лист
    if include_title_page:
        _add_title_page(doc, state)

    # Оглавление
    if include_toc:
        _add_table_of_contents(doc, state)

    # Разделы
    _add_sections(doc, state)

    # Сохранение
    if output_path is None:
        object_name = (state.card or {}).get("object_name", "ПЗ")
        # Очищаем имя от недопустимых символов
        safe_name = "".join(c for c in object_name if c.isalnum() or c in " _-")
        safe_name = safe_name.strip() or "Пояснительная_записка"
        output_dir = Path(__file__).resolve().parent.parent / "output"
        output_dir.mkdir(exist_ok=True)
        output_path = str(output_dir / f"{safe_name}_ОТР.docx")

    doc.save(output_path)
    return output_path


# ============================================================
# ТЕСТ
# ============================================================

if __name__ == "__main__":
    from pipeline._04_orchestrator import PipelineState

    # Создаём тестовое состояние
    state = PipelineState()
    state.card = {
        "object_name": 'ПС 110/10 кВ "Стойленская"',
        "voltage_hv": "110 кВ",
        "voltage_lv": "10 кВ",
        "scheme_hv": "4Н",
    }
    state.sections = [
        {
            "section_num": "1",
            "section_title": "Общие сведения",
            "text": (
                "В настоящем разделе приведены основные технические решения "
                "по релейной защите и автоматике подстанции 110/10 кВ.\n\n"
                "Проектирование выполняется в соответствии с действующими "
                "нормативными документами. Состав защит определяется "
                "требованиями ПУЭ и СТО."
            ),
            "sources": [
                {"num": 1, "source": "ПУЭ гл. 3.2", "section": "3.2", "collection": "gost"},
            ],
            "is_fallback": False,
        },
        {
            "section_num": "2",
            "section_title": "Защита трансформатора 25 МВА",
            "text": (
                "Для защиты силового трансформатора 25 МВА предусматриваются "
                "следующие типы защит:\n\n"
                "ОСНОВНЫЕ ЗАЩИТЫ\n"
                "— дифференциальная токовая защита (ДЗТ);\n"
                "— газовая защита.\n\n"
                "РЕЗЕРВНЫЕ ЗАЩИТЫ\n"
                "— максимальная токовая защита (МТЗ) на стороне ВН;\n"
                "— МТЗ на стороне НН."
            ),
            "sources": [
                {"num": 1, "source": "СТО 34.01-21-005-2019", "section": "5.3", "collection": "gost"},
                {"num": 2, "source": "Проект ПС 110 кВ Ивановская", "section": "4.2", "collection": "projects"},
            ],
            "is_fallback": False,
        },
        {
            "section_num": "3",
            "section_title": "Защита БСК",
            "text": "Параметры и состав защит определяются на этапе РД по согласованию с заводом-изготовителем.",
            "sources": [],
            "is_fallback": True,
        },
    ]
    state.status = "done"

    print("=" * 60)
    print("ТЕСТ RZA_Docx_Assembler")
    print("=" * 60)

    output_path = assemble_docx(state)
    print(f"✅ Документ сохранён: {output_path}")
